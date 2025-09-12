import os   #Módulo para navegar por los directorio de windows
import json #Móduo para generar los archivos de formato .json
from docx import Document   #Módulo para trabajar y leer documentos de fromato .docx (word)
import typing
import logging
import sys

logger = logging.getLogger(__name__)  #Manejo de errores y logs durante la ejecución del programa
logging.basicConfig(filename="Debugging_logs.log", encoding='utf-8', level=logging.DEBUG)

class DocumentProcessor:  #Clase para procesar documentos de tipo docx
    def __init__(self, path: str): #Constructor que recibe como atributos la ruta de los archivos .docx, y declara como atributos globales el nombre del archivo a procesar y el nombre de la entidad que será extraido en la ejecución del programa
        self.path = path
        self.filename = os.path.basename(path)
        self.entity_name = None  # se extraerá del encabezado de una solo página, tomando en cuenta que es el único encabezado y se encontrará siempre ahí el nombre de la entidad

    def extract_entity_from_header(self, doc): # Método para extraer el nombre de la entidad desde la tabla del encabezado
        for section in doc.sections: # Bucle for principal para iterar cada una de las secciones del documento donde se encuetra el encabezado
            header = section.header # Cada seccion contiene un encabezado
            for table in header.tables:# Se iterará la tabla de cada encabezado, en este caso siempre será uncamente uno
                for row in table.rows:# De cada tabla, se iterará cada fila
                    for i, cell in enumerate(row.cells):# Se itera y se enumera cada celda de la columna con el método .cells del objeto "row"
                        if "Entidad Fiscalizada:" in cell.text: #Si encuentra el texto "Entidad fisca", en la celda iterada, recorrera una celda mas para extraer el valor de la siguiente celda que en este caso siempre será el nombre de la entidad fiscalizada
                            # Tomamos la celda siguiente mientras no exceda el límite de celdas de la fila
                            if i + 1 < len(row.cells):
                                return row.cells[i + 1].text.strip()
        return "Entidad no encontrada"

    def docx_processor(self):  #método para iniciar con el proceso de extraccion de la informacion de los documentos
        doc = Document(self.path) # se define variable que recibirá la ruta del archivo a procesar

        # Extraer entidad desde el encabezado
        self.entity_name = self.extract_entity_from_header(doc)  #Se llama al metodo estático de la clase para extraer el nombre de la entidad através del encabezado

        # Concatenar todo el texto de párrafos y tablas
        texts = []  #Este arreglo se encarga de concatenar todo el texto del archivo completo
        for p in doc.paragraphs:  #procesamiento y concatenación de párrafos, en este caso no existen parrafos ya que toda la informacion se encuentra dentro de tablas.
            if p.text.strip(): # si existe un texto, lo concatenara en el arreglo
                texts.append(p.text.strip())
        for table in doc.tables: # bucle encargado de iterar todas las tablas que se vayan encontrando
            for row in table.rows:# Dentro de cada tabla, se iterarán todas sus filas
                for cell in row.cells:# Dentro de cada fila, se iteran todas sus celdas
                    if cell.text.strip(): # Si existe algun texto, lo concatenará al arreglo
                        texts.append(cell.text.strip())

        total_text = "\n".join(texts) #Funcion .join para concatenar y separar por medio de in \n cada elemento del arreglo, dando como resultado asi un solo texto junto, pero serparado por \n

        # Separar observaciones por “OBSERVACIÓN NO.”
        obs_chunks = total_text.split("OBSERVACIÓN NO.") #Dada la cadena de texto completa, se hará una división por "chunks" cada nuevo chunk empieza donde se encuentra la palabra "OBSERVACIÓN NO." esto con el objetivo de construir un arreglo en el cada elemento identificado a base de los chunks, represnte cada observación o cada tabla

        observations_fetched = [] #Arrelgo que alamcenará diccionarios de cada observación o tabla encontrada
        for chunk in obs_chunks[1:]:  # Se itera cada chunk (observación) con sus cadenas de texto donde se continene los datos importantes 
            lines = chunk.strip().split("\n") #Desde la lina de codigo No.41 se estableció que cada seccion de las tablas o párrafos se depararán con un \n al momenot de concatenar todos los elementos del arreglo total_text, en esta linea se separan cada valor de las celdas usando el \n y los mete como elementos de este arreglo
            no_observation = int(lines[0].split()[0]) #extraccion y alamcenarmiento del número de observación

            chapter = theme = subtheme = "" #inicializacion de las variables que contendrán cada dato importante a extraer
            observation_lines = [] #arreglo que almacena la descripción de la observación
            legal = None #de momento no es necesrio extraer el fundamento legal

            for line in lines[1:]: #Se itera el arreglo que contiene como elemnto cada lina del chunk
                if line.startswith("Capítulo:"): #condicion para revisar si el texto de la lina contiene el string para guardar el dato 
                    chapter = line.replace("Capítulo:", "").strip()
                elif line.startswith("Tema:"):
                    theme = line.replace("Tema:", "").strip()
                elif line.startswith("Subtema:"):
                    subtheme = line.replace("Subtema:", "").strip()
                elif line.startswith("MARCO LEGAL"):
                    break
                else:
                    observation_lines.append(line.strip())

            observation_text = " ".join(observation_lines).strip() #Concatenacion de todas las lineas que conforman la observacion

            observations_fetched.append({ #funcion append para agregar el diccionario con todos los datos extraídos al arreglo 
                "no_observation": no_observation,
                "chapter": chapter,
                "theme": theme,
                "subtheme": subtheme,
                "observation": observation_text,
                "legal": legal
            })

        result = { #atributos principales de cada docuemnto en el formato json
            "filename": self.filename,
            "entity": self.entity_name,
            "observations_fetched": observations_fetched
        }
        return result #retorna el json final


def process_directory(input_dir, output_file = "output/output.json") -> bool:
    all_documents = {"documents": []}
    for file in os.listdir(input_dir):
        if not file.endswith(".docx"):
            logging.error(f"Archivo inválido (no .docx): {file}")
            continue  
        path = os.path.join(input_dir, file)
        logging.info(f"Procesando archivo: {file}")
        processor = DocumentProcessor(path)
        result = processor.docx_processor()
        all_documents["documents"].append(result)
        logging.info(f"✅ Archivo procesado: {file}")
    return True
    


    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(all_documents, f, ensure_ascii=False, indent=4)
        
    return all_documents


#result = process_directory("C:/Users/esfe0/Documents/ESFE/semáforo_observaciones/test_files","../output/output.json")



